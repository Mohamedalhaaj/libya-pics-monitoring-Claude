"""Tests for utils.report_eval against a synthetic PICS report.

Runnable directly (`python tests/test_report_eval.py`) or under pytest.
"""

from __future__ import annotations

import sys
import tempfile
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from docx import Document  # noqa: E402

from utils.report_eval import (  # noqa: E402
    canonical_source,
    coverage_recall,
    parse_pics_report,
    report_metrics,
)


def _write_report(path: Path, bullets_by_section: dict[str, list[str]], title: str) -> None:
    doc = Document()
    doc.add_paragraph(title)
    for section, bullets in bullets_by_section.items():
        doc.add_paragraph(section)
        for line in bullets:
            doc.add_paragraph(line, style="List Bullet")
    doc.save(str(path))


GOLD = {
    "United Nations": [
        "UN envoy warns window for action is narrowing – UN News / Libya Observer",
        "[Mufti] Gharyani issues a fatwa rejecting the outcomes – Al Marsad (Arabic)",
    ],
    "Politics": [
        "Libya's three councils agree to hold elections before 2027 – Libya Observer / Anadolu Agency",
        "47 HoR members voice support for Boulos initiative – Al Wasat (Arabic)",
    ],
}

# Candidate: same structure, but one headline untranslated and one outlet missing.
CANDIDATE = {
    "United Nations": [
        "UN envoy warns window for action is narrowing – UN News",
        "غريان يصدر فتوى برفض المخرجات – Al Marsad (Arabic)",
    ],
    "Politics": [
        "Libya's three councils agree to hold elections before 2027 – Libya Observer",
    ],
}


def test_parse_and_metrics():
    with tempfile.TemporaryDirectory() as tmp:
        path = Path(tmp) / "gold.docx"
        _write_report(path, GOLD, "Libya News Headlines – 18-21 June")
        report = parse_pics_report(path)
        assert report.title.startswith("Libya News Headlines")
        assert len(report.bullets) == 4
        assert report.sections == ["United Nations", "Politics"]
        m = report_metrics(report)
        assert m["title_ok"] and m["sections_ordered"]
        assert m["has_source_ratio"] == 1.0
        assert m["noise_bullets"] == 0
        assert m["duplicate_bullets"] == 0 and m["distinct_ratio"] == 1.0
        assert m["english_ratio"] == 1.0  # all gold bullets are English
        assert 0.0 < m["role_prefix_ratio"] <= 0.5  # one "[Mufti] ..." bullet
        assert m["multi_source_ratio"] == 0.5  # two of four cite >=2 outlets


def test_coverage_recall_and_english_detection():
    with tempfile.TemporaryDirectory() as tmp:
        gpath, cpath = Path(tmp) / "g.docx", Path(tmp) / "c.docx"
        _write_report(gpath, GOLD, "Libya News Headlines – 18-21 June")
        _write_report(cpath, CANDIDATE, "Libya News Headlines – 18-21 June")
        gold = parse_pics_report(gpath)
        cand = parse_pics_report(cpath)

        # Candidate has 3 bullets, one in Arabic -> english_ratio 2/3.
        mc = report_metrics(cand)
        assert mc["n_bullets"] == 3
        assert abs(mc["english_ratio"] - 2 / 3) < 1e-6

        cov = coverage_recall(cand, gold)
        gold_outlets = gold.source_names()
        # Gold cites: un news, libya observer, al marsad, anadolu agency, al wasat (5).
        assert cov["gold_sources"] == len(gold_outlets) == 5
        # Candidate cites un news, al marsad, libya observer (3) -> recall 3/5.
        assert cov["matched_sources"] == 3
        assert abs(cov["source_recall"] - 0.6) < 1e-6
        assert "anadolu agency" in cov["missing_sources"]


def test_duplicate_bullets_detected():
    dup = {"United Nations": [
        "UN envoy warns the window for action is narrowing – UN News",
        "UN envoy warns the window for action is narrowing. – Libya Observer",  # same, punctuation aside
        "Tetteh briefs the Security Council on Libya – Libya Review",
    ]}
    with tempfile.TemporaryDirectory() as tmp:
        path = Path(tmp) / "d.docx"
        _write_report(path, dup, "Libya News Headlines – 18-21 June")
        m = report_metrics(parse_pics_report(path))
        assert m["n_bullets"] == 3
        assert m["duplicate_bullets"] == 1  # first two normalise equal
        assert abs(m["distinct_ratio"] - 2 / 3) < 1e-6


def test_canonical_source_matching():
    # Naming variants that should compare equal across reports.
    assert canonical_source("Asharq Al-Awsat") == canonical_source("Asharq Al Awsat")
    assert canonical_source("Libya Herald (English)") == canonical_source("Libya Herald")
    assert canonical_source("The New Arab (English)") == canonical_source("New Arab")


def test_ceiling_recall_vs_monitored():
    with tempfile.TemporaryDirectory() as tmp:
        gpath, cpath = Path(tmp) / "g.docx", Path(tmp) / "c.docx"
        _write_report(gpath, GOLD, "Libya News Headlines – 18-21 June")
        _write_report(cpath, CANDIDATE, "Libya News Headlines – 18-21 June")
        gold = parse_pics_report(gpath)
        cand = parse_pics_report(cpath)

        # Gold cites 5 outlets; we only "monitor" 3 of them. The candidate cites
        # all 3 monitored ones -> ceiling recall 100% even though raw recall 60%.
        monitored = {"UN News", "Al Marsad", "Libya Observer"}
        cov = coverage_recall(cand, gold, monitored=monitored)
        assert cov["achievable_sources"] == 3
        assert cov["achievable_matched"] == 3
        assert cov["unmonitored_gold"] == 2  # anadolu agency, al wasat
        assert abs(cov["ceiling_recall"] - 1.0) < 1e-6
        assert abs(cov["source_recall"] - 0.6) < 1e-6
        assert cov["missing_monitored"] == []


if __name__ == "__main__":
    tests = [v for k, v in sorted(globals().items()) if k.startswith("test_") and callable(v)]
    for test in tests:
        test()
        print(f"PASS {test.__name__}")
    print(f"\nAll {len(tests)} report-eval tests passed.")
