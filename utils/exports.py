from __future__ import annotations

import csv
import re
import unicodedata
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.opc.constants import RELATIONSHIP_TYPE as REL_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from utils import taxonomy
from utils.models import (
    Article,
    HeadlineSource,
    ReportHeadline,
    ReportSection,
    ReportSubsection,
    SourceVerification,
    StructuredReport,
)


ARTICLE_FIELDS = [
    "source_id",
    "source_name",
    "language",
    "country_focus",
    "title",
    "url",
    "published_at",
    "summary",
    "section",
    "matched_keywords",
]

VERIFICATION_FIELDS = [
    "source_id",
    "source_name",
    "url",
    "status",
    "articles_found",
    "error",
    "checked_at",
]


def ensure_output_dir(path: str | Path) -> Path:
    output_dir = Path(path)
    output_dir.mkdir(parents=True, exist_ok=True)
    return output_dir


def write_articles_csv(articles: list[Article], path: str | Path) -> None:
    with Path(path).open("w", encoding="utf-8-sig", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=ARTICLE_FIELDS)
        writer.writeheader()
        for article in articles:
            writer.writerow(article.to_row())


def write_verification_csv(verifications: list[SourceVerification], path: str | Path) -> None:
    with Path(path).open("w", encoding="utf-8-sig", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=VERIFICATION_FIELDS)
        writer.writeheader()
        for verification in verifications:
            writer.writerow(verification.to_row())


# --------------------------------------------------------------------------- #
# Mechanical fallback: shape scraped articles into a StructuredReport without  #
# the Claude enrichment layer. Used when enrichment is unavailable/disabled.   #
# --------------------------------------------------------------------------- #

# Lightweight keyword → section routing for the fallback path. The enrichment
# layer does this far better; this only needs to be a reasonable default.
_SECTION_KEYWORDS: list[tuple[str, tuple[str, ...]]] = [
    (
        "United Nations",
        ("unsmil", "united nations", "srsg", "tetteh", "security council",
         "البعثة الأممية", "الأمم المتحدة", "مجلس الأمن"),
    ),
    (
        "Human Rights & Rule of Law",
        ("migrant", "migration", "refugee", "human rights", "trafficking",
         "هجرة", "مهاجر", "حقوق الإنسان", "لاجئ"),
    ),
    (
        "Economy",
        ("oil", "central bank", "dinar", "dollar", "fuel", "economy", "budget",
         "نفط", "المصرف المركزي", "الدينار", "الدولار", "وقود", "اقتصاد"),
    ),
    (
        "Military & Security",
        ("army", "security", "military", "clash", "weapon", "arrest", "border",
         "الجيش", "الأمن", "عسكري", "اشتباك", "سلاح", "الحدود"),
    ),
    (
        "Environment",
        ("flood", "weather", "earthquake", "climate", "rain", "drought",
         "فيضان", "طقس", "زلزال", "مناخ", "أمطار"),
    ),
    (
        "Regional & International",
        ("tunisia", "egypt", "italy", "turkey", "european", "foreign",
         "تونس", "مصر", "إيطاليا", "تركيا", "خارجية"),
    ),
    (
        "Politics",
        ("government", "parliament", "election", "political", "hor", "gnu",
         "حكومة", "البرلمان", "انتخابات", "سياسي", "مجلس النواب", "مجلس الدولة"),
    ),
]


def _normalize(text: str) -> str:
    return unicodedata.normalize("NFKC", text).casefold()


def _classify_section(article: Article) -> str:
    haystack = _normalize(f"{article.title} {article.summary} {article.section}")
    for section, keywords in _SECTION_KEYWORDS:
        if any(_normalize(keyword) in haystack for keyword in keywords):
            return section
    return "Varieties"


def _title_key(title: str) -> str:
    return re.sub(r"\s+", " ", _normalize(re.sub(r"[^\w\s]", " ", title))).strip()


def build_fallback_report(articles: list[Article], report_date: str) -> StructuredReport:
    """Group scraped articles into the report taxonomy without the LLM.

    Articles with near-identical titles are merged into a single headline that
    cites every reporting source. This produces the correct report *shape*;
    headlines remain in their original (often Arabic) wording.
    """
    # bucket: section -> title_key -> ReportHeadline
    buckets: dict[str, dict[str, ReportHeadline]] = {}
    for article in articles:
        section = _classify_section(article)
        key = _title_key(article.title) or article.url or article.title
        headline_map = buckets.setdefault(section, {})
        source = HeadlineSource(name=article.source_name, language=article.language, url=article.url)
        if key in headline_map:
            existing = headline_map[key]
            if not any(s.name == source.name for s in existing.sources):
                existing.sources.append(source)
        else:
            headline_map[key] = ReportHeadline(text=article.title, sources=[source])

    sections: list[ReportSection] = []
    for section_name in taxonomy.SECTION_ORDER:
        headline_map = buckets.get(section_name)
        if not headline_map:
            continue
        subsection = ReportSubsection(name="", headlines=list(headline_map.values()))
        sections.append(ReportSection(name=section_name, subsections=[subsection]))

    return StructuredReport(report_date=report_date, sections=sections, is_draft=True)


# --------------------------------------------------------------------------- #
# Word rendering                                                              #
# --------------------------------------------------------------------------- #

def _add_hyperlink(paragraph, url: str, text: str) -> None:
    part = paragraph.part
    r_id = part.relate_to(url, REL_TYPE.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run = OxmlElement("w:r")
    run_props = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    run_props.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_props.append(underline)
    run.append(run_props)

    text_el = OxmlElement("w:t")
    text_el.text = text
    text_el.set(qn("xml:space"), "preserve")
    run.append(text_el)

    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _render_headline(document: Document, headline: ReportHeadline) -> None:
    # "List Bullet" gives the bulleted layout the PICS house style uses.
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.add_run(headline.text)
    if headline.sources:
        paragraph.add_run(" – ")
        for index, source in enumerate(headline.sources):
            if index:
                paragraph.add_run(" / ")
            label = source.render()
            if source.url:
                _add_hyperlink(paragraph, source.url, label)
            else:
                paragraph.add_run(label)


def _add_section_heading(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(14)


def _add_subsection_heading(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = True


def write_word_report(
    report: StructuredReport,
    verifications: list[SourceVerification],
    path: str | Path,
) -> None:
    """Render a StructuredReport in the UNSMIL/PICS Libya News Headlines format."""
    document = Document()

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run(f"Libya News Headlines – {report.report_date}")
    title_run.bold = True
    title_run.font.size = Pt(16)

    if report.is_draft:
        banner = document.add_paragraph()
        banner.alignment = WD_ALIGN_PARAGRAPH.CENTER
        banner_run = banner.add_run(
            "DRAFT — mechanical layout: headlines are NOT translated or "
            "editorially deduplicated. Set ANTHROPIC_API_KEY and re-run "
            "(without --no-enrich) to produce the final report."
        )
        banner_run.bold = True
        banner_run.italic = True
        banner_run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    if not report.sections:
        document.add_paragraph(
            "No matching Libya-related articles were collected for the selected date range."
        )

    for section in report.sections:
        _add_section_heading(document, section.name)
        for subsection in section.subsections:
            if subsection.name:
                _add_subsection_heading(document, subsection.name)
            for headline in subsection.headlines:
                _render_headline(document, headline)

    # Source verification appendix (operational record, kept out of the body).
    document.add_page_break()
    appendix = document.add_paragraph()
    appendix_run = appendix.add_run("Source Verification")
    appendix_run.bold = True
    appendix_run.font.size = Pt(14)

    table = document.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    for index, header in enumerate(["Source", "URL", "Status", "Articles", "Error"]):
        table.rows[0].cells[index].text = header
    for verification in verifications:
        cells = table.add_row().cells
        cells[0].text = verification.source_name
        cells[1].text = verification.url or ""
        cells[2].text = verification.status
        cells[3].text = str(verification.articles_found)
        cells[4].text = verification.error or ""

    disclaimer = document.add_paragraph()
    disclaimer_run = disclaimer.add_run(taxonomy.DISCLAIMER)
    disclaimer_run.italic = True
    disclaimer_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    document.save(path)
